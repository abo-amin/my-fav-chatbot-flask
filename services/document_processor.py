"""
Document Processor - Handles text extraction from various file formats
Supports: PDF, DOCX, CSV, TXT, XLSX
"""
import os
from pathlib import Path
from typing import List, Dict, Optional
import config


class DocumentProcessor:
    """Process and extract text from documents"""
    
    def __init__(self):
        self.chunk_size = config.CHUNK_SIZE
        self.chunk_overlap = config.CHUNK_OVERLAP
    
    def process_file(self, file_path: str) -> Dict:
        """Process a file and extract its text content"""
        path = Path(file_path)
        file_type = path.suffix.lower().lstrip('.')
        
        if file_type == 'pdf':
            text = self._extract_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            text = self._extract_docx(file_path)
        elif file_type in ['csv']:
            text = self._extract_csv(file_path)
        elif file_type in ['xlsx', 'xls']:
            text = self._extract_excel(file_path)
        elif file_type == 'txt':
            text = self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Create chunks
        chunks = self._create_chunks(text, path.name)
        
        return {
            'filename': path.name,
            'file_type': file_type,
            'text': text,
            'chunks': chunks,
            'chunk_count': len(chunks)
        }
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            
            text_parts = []
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}")
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")
    
    def _extract_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            
            # Create a readable text representation
            text_parts = []
            text_parts.append(f"CSV File with {len(df)} rows and {len(df.columns)} columns")
            text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
            text_parts.append("\nData:")
            
            # Convert each row to readable text
            for idx, row in df.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                text_parts.append(f"Row {idx + 1}: {row_text}")
            
            return "\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Error extracting CSV: {str(e)}")
    
    def _extract_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_parts = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text_parts.append(f"\n[Sheet: {sheet_name}]")
                text_parts.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
                text_parts.append(f"Columns: {', '.join(df.columns.astype(str).tolist())}")
                
                for idx, row in df.iterrows():
                    row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                    text_parts.append(f"Row {idx + 1}: {row_text}")
            
            return "\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Error extracting Excel: {str(e)}")
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error extracting TXT: {str(e)}")
    
    def _create_chunks(self, text: str, filename: str) -> List[Dict]:
        """Split text into overlapping chunks"""
        if not text.strip():
            return []
        
        # Split by paragraphs first
        paragraphs = text.split('\n\n')
        
        chunks = []
        current_chunk = ""
        chunk_index = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # If adding this paragraph exceeds chunk size, save current and start new
            if len(current_chunk) + len(para) > self.chunk_size * 4:  # Approximate char count
                if current_chunk:
                    chunks.append({
                        'index': chunk_index,
                        'content': current_chunk.strip(),
                        'metadata': f"Source: {filename}, Chunk: {chunk_index + 1}"
                    })
                    chunk_index += 1
                    
                    # Keep overlap
                    words = current_chunk.split()
                    overlap_words = words[-self.chunk_overlap:] if len(words) > self.chunk_overlap else []
                    current_chunk = ' '.join(overlap_words) + '\n\n' + para
                else:
                    current_chunk = para
            else:
                current_chunk = current_chunk + '\n\n' + para if current_chunk else para
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append({
                'index': chunk_index,
                'content': current_chunk.strip(),
                'metadata': f"Source: {filename}, Chunk: {chunk_index + 1}"
            })
        
        return chunks
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return list(config.ALLOWED_EXTENSIONS)


# Singleton instance
document_processor = DocumentProcessor()
